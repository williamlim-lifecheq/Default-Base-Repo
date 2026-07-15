#!/usr/bin/env python3
"""
Generates the next monthly contractor invoice from the Lifecheq docx
template (Sample_Contractor_Invoice_Template), continuing the invoice
number sequence from whatever invoice is most recently on file (docx or
legacy xlsx).

Usage:
    python3 generate_invoice.py                  # next calendar month after the latest invoice on file
    python3 generate_invoice.py --month 8 --year 2026
    python3 generate_invoice.py --fee 12000       # override the base consultation fee
    python3 generate_invoice.py --invoice-number LC010   # override auto-increment

Run it from inside the "Lifecheq Admin" folder (or pass --dir).
"""
import argparse
import calendar
import datetime
import re
from pathlib import Path

import docx
import openpyxl

TEMPLATE_NAME = "Sample_Contractor_Invoice_Template (1) (3).docx"

DOCX_RE = re.compile(r"^Invoicing - Lim Kok Leong - (?P<month>[A-Za-z]+) (?P<year>\d{4})\.docx$")
XLSX_RE = re.compile(r"^Invoicing - Lim Kok Leong - (?P<month>[A-Za-z]+) (?P<year>\d{4})\.xlsx$")

CONTRACTOR_NAME = "LIM Kok Leong"
CONTRACTOR_ADDRESS_LINE1 = "Unit C-19-07, Nidoz Residences"
CONTRACTOR_ADDRESS_LINE2 = "No. 22A, Jln 2/125, 57100 Kuala Lumpur"
CONTRACTOR_EMAIL = "william.lim@lifecheq.co.za"

BANK_NAME = "RHB Bank Berhad"
BANK_ACCOUNT_NUMBER = "21407100150515"
BANK_SWIFT = "RHBBMYKLXXX"
BANK_ADDRESS = "Two Tower Two & Three, RHB Centre, Jalan Tun Razak, 50400 Kuala Lumpur"

CURRENCY = "USD"


def fmt_amount(amount):
    return f"{CURRENCY} {amount:,.2f}"


def find_last_invoice(directory):
    """Return (year, month, invoice_number) for the most recent invoice on file, docx or legacy xlsx."""
    candidates = []
    for path in directory.glob("Invoicing - Lim Kok Leong - *.docx"):
        m = DOCX_RE.match(path.name)
        if not m:
            continue
        year = int(m.group("year"))
        month = datetime.datetime.strptime(m.group("month"), "%B").month
        d = docx.Document(path)
        inv_cell_text = d.tables[0].rows[0].cells[1].paragraphs[1].text
        num_match = re.search(r"Invoice #:\s*(\S+)", inv_cell_text)
        if num_match:
            candidates.append((year, month, num_match.group(1), path))

    for path in directory.glob("Invoicing - Lim Kok Leong - *.xlsx"):
        m = XLSX_RE.match(path.name)
        if not m:
            continue
        year = int(m.group("year"))
        month = datetime.datetime.strptime(m.group("month"), "%B").month
        wb = openpyxl.load_workbook(path, data_only=False)
        invoice_number = wb.active["B15"].value
        candidates.append((year, month, invoice_number, path))

    if not candidates:
        return None
    candidates.sort(key=lambda c: (c[0], c[1]))
    return candidates[-1]


def next_invoice_number(current):
    m = re.match(r"^([A-Za-z]+)(\d+)$", current)
    if not m:
        raise ValueError(f"Can't parse invoice number '{current}'")
    prefix, digits = m.group(1), m.group(2)
    return f"{prefix}{int(digits) + 1:0{len(digits)}d}"


def set_run_text(paragraph, index, new_text):
    paragraph.runs[index].text = new_text


def replace_bracket(text, replacement):
    return re.sub(r"\[.*?\]", replacement, text, count=1)


def fill_template(template_path, out_path, invoice_number, invoice_date, due_date, description, qty, fee):
    d = docx.Document(template_path)

    header_left, header_right = d.tables[0].rows[0].cells
    set_run_text(header_left.paragraphs[0], 0, CONTRACTOR_NAME)
    set_run_text(header_left.paragraphs[1], 0, CONTRACTOR_ADDRESS_LINE1)
    set_run_text(header_left.paragraphs[2], 0, CONTRACTOR_ADDRESS_LINE2)
    set_run_text(header_left.paragraphs[3], 0, CONTRACTOR_EMAIL)
    set_run_text(header_left.paragraphs[4], 0, "Not VAT registered")

    inv_para = header_right.paragraphs[1]
    inv_para.runs[0].text = replace_bracket(inv_para.runs[0].text, invoice_number)

    dates_table = d.tables[1].rows[0].cells[1].tables[0]
    set_run_text(dates_table.rows[0].cells[1].paragraphs[0], 0, invoice_date.strftime("%d/%m/%Y"))
    set_run_text(dates_table.rows[1].cells[1].paragraphs[0], 0, due_date.strftime("%d/%m/%Y"))

    services = d.tables[2]
    set_run_text(services.rows[1].cells[0].paragraphs[0], 0, description)
    set_run_text(services.rows[1].cells[1].paragraphs[0], 0, str(qty))
    set_run_text(services.rows[1].cells[2].paragraphs[0], 0, fmt_amount(fee))
    for row in services.rows[2:]:
        set_run_text(row.cells[0].paragraphs[0], 0, "")
        set_run_text(row.cells[1].paragraphs[0], 0, "")
        set_run_text(row.cells[2].paragraphs[0], 0, "")

    totals = d.tables[3]
    set_run_text(totals.rows[0].cells[2].paragraphs[0], 0, fmt_amount(fee))  # Subtotal
    set_run_text(totals.rows[1].cells[2].paragraphs[0], 0, "N/A")            # VAT (not applicable)
    set_run_text(totals.rows[2].cells[2].paragraphs[0], 0, fmt_amount(fee))  # Total due

    banking = d.tables[4]
    set_run_text(banking.rows[0].cells[1].paragraphs[0], 0, CONTRACTOR_NAME)
    set_run_text(banking.rows[1].cells[1].paragraphs[0], 0, BANK_NAME)
    set_run_text(banking.rows[2].cells[1].paragraphs[0], 0, BANK_ACCOUNT_NUMBER)
    set_run_text(banking.rows[3].cells[1].paragraphs[0], 0, BANK_SWIFT)
    addr_para = banking.rows[4].cells[1].paragraphs[0]
    addr_para.runs[0].text = BANK_ADDRESS
    for extra_run in addr_para.runs[1:]:
        extra_run.text = ""

    notes_para = d.paragraphs[7]
    notes_para.runs[0].text = "Payment due by the due date specified above."

    d.save(out_path)


def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--dir", default=".", help="Folder containing the invoice files (default: current dir)")
    parser.add_argument("--month", type=int, help="Target month (1-12). Default: month after the latest invoice on file")
    parser.add_argument("--year", type=int, help="Target year. Default: inferred alongside --month")
    parser.add_argument("--invoice-date-day", type=int, default=15, help="Day of month for the invoice date (default: 15)")
    parser.add_argument("--fee", type=float, default=11000, help="Base Consultation Services fee (default: 11000)")
    parser.add_argument("--qty", default="1", help="Qty / Hrs value for the base fee line (default: 1)")
    parser.add_argument("--description", default="Consultation Services", help="Description for the base fee line")
    parser.add_argument("--invoice-number", help="Override the auto-incremented invoice number")
    parser.add_argument("--dry-run", action="store_true", help="Print what would be generated without writing a file")
    args = parser.parse_args()

    directory = Path(args.dir).resolve()
    template_path = directory / TEMPLATE_NAME
    if not template_path.exists():
        raise SystemExit(f"Template not found: {template_path}")

    last = find_last_invoice(directory)
    if not last:
        raise SystemExit(f"No existing 'Invoicing - Lim Kok Leong - <Month> <Year>' files found in {directory}")
    latest_year, latest_month, latest_number, latest_path = last

    if args.month:
        target_month = args.month
        target_year = args.year or (latest_year if target_month > latest_month else latest_year + 1)
    else:
        if latest_month == 12:
            target_month, target_year = 1, latest_year + 1
        else:
            target_month, target_year = latest_month + 1, latest_year

    target_month_name = calendar.month_name[target_month]
    new_filename = f"Invoicing - Lim Kok Leong - {target_month_name} {target_year}.docx"
    new_path = directory / new_filename
    if new_path.exists():
        raise SystemExit(f"{new_filename} already exists — refusing to overwrite it.")

    new_number = args.invoice_number or next_invoice_number(latest_number)
    invoice_date = datetime.date(target_year, target_month, args.invoice_date_day)
    last_day = calendar.monthrange(target_year, target_month)[1]
    due_date = datetime.date(target_year, target_month, last_day)

    print(f"Latest invoice on file : {latest_path.name} (invoice {latest_number})")
    print(f"Generating             : {new_filename}")
    print(f"  Invoice number       : {new_number}")
    print(f"  Invoice date         : {invoice_date.strftime('%d/%m/%Y')}")
    print(f"  Due date             : {due_date.strftime('%d/%m/%Y')}")
    print(f"  Base fee             : {fmt_amount(args.fee)}")

    if args.dry_run:
        print("Dry run — no file written.")
        return

    fill_template(template_path, new_path, new_number, invoice_date, due_date, args.description, args.qty, args.fee)
    print(f"Wrote {new_path}")
    print("Note: Subtotal/VAT/Total are plain text, not live formulas — if you add claim lines, "
          "update the Subtotal and Total Due amounts by hand before sending it to the accountant.")


if __name__ == "__main__":
    main()
